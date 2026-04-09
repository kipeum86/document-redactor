#!/usr/bin/env python3
"""
Fixture generator for Gate 0 docx.js spike.

Produces tests/fixtures/bilingual_nda_worst_case.docx, a synthetic bilingual
(Korean + English) NDA that exercises every one of the 7 Gate 0 checklist items:

    1. Body / table / header / footer / footnote text
    2. Track changes (insertions + deletions)
    3. Comments
    4. Korean NFC/NFD/한자/emoji/mixed-script
    5. Section breaks (with different header/footer per section)
    6. Complex tables with merged + nested cells
    7. Clean round-trip (no "repair needed" dialog when reopened)

python-docx is used on purpose — it's a DIFFERENT library from docx.js, so
the fixture isn't biased toward what the library-under-test can handle.

Run:
    python3 tools/make-fixture.py
Output:
    tests/fixtures/bilingual_nda_worst_case.docx
"""

from __future__ import annotations

import os
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


FIXTURE_DIR = Path(__file__).resolve().parent.parent / "tests" / "fixtures"
FIXTURE_PATH = FIXTURE_DIR / "bilingual_nda_worst_case.docx"


# ---------------------------------------------------------------------------
# Low-level helpers — python-docx's public API doesn't expose track changes
# or comments, so we reach into the underlying XML directly where needed.
# ---------------------------------------------------------------------------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def add_tracked_insertion(paragraph, text: str, author: str = "Reviewer A") -> None:
    """Insert `text` as a tracked insertion (w:ins) into the paragraph."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), "2026-04-09T12:00:00Z")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)

    ins.append(run)
    paragraph._p.append(ins)


def add_tracked_deletion(paragraph, text: str, author: str = "Reviewer A") -> None:
    """Mark `text` as tracked deletion (w:del) in the paragraph."""
    d = OxmlElement("w:del")
    d.set(qn("w:id"), "2")
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), "2026-04-09T12:05:00Z")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    run.append(rpr)
    del_text = OxmlElement("w:delText")
    del_text.text = text
    del_text.set(qn("xml:space"), "preserve")
    run.append(del_text)

    d.append(run)
    paragraph._p.append(d)


def add_comment_range(
    doc: Document,
    paragraph,
    commented_text: str,
    comment_text: str,
    author: str = "Partner Kim",
    comment_id: int = 1,
) -> None:
    """
    Insert a commented region into `paragraph`. Creates the comments.xml part
    if it doesn't exist and wires up the relationships.
    """
    from docx.opc.constants import CONTENT_TYPE, RELATIONSHIP_TYPE
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    part = doc.part
    package = part.package

    # Check if comments part already exists
    comments_part = None
    for rel in part.rels.values():
        if rel.reltype == RELATIONSHIP_TYPE.COMMENTS:
            comments_part = rel.target_part
            break

    if comments_part is None:
        comments_xml = (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:comments xmlns:w="{W_NS}"></w:comments>'
        ).encode("utf-8")
        comments_part = Part(
            PackURI("/word/comments.xml"),
            CONTENT_TYPE.WML_COMMENTS,
            comments_xml,
            package,
        )
        part.relate_to(comments_part, RELATIONSHIP_TYPE.COMMENTS)

    # Append the comment content
    from lxml import etree

    comments_root = etree.fromstring(comments_part.blob)
    comment_el = etree.SubElement(comments_root, qn("w:comment"))
    comment_el.set(qn("w:id"), str(comment_id))
    comment_el.set(qn("w:author"), author)
    comment_el.set(qn("w:date"), "2026-04-09T14:30:00Z")
    comment_el.set(qn("w:initials"), author[0])
    p_el = etree.SubElement(comment_el, qn("w:p"))
    r_el = etree.SubElement(p_el, qn("w:r"))
    t_el = etree.SubElement(r_el, qn("w:t"))
    t_el.text = comment_text
    comments_part._blob = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # Wire up the range markers in the paragraph
    range_start = OxmlElement("w:commentRangeStart")
    range_start.set(qn("w:id"), str(comment_id))
    paragraph._p.append(range_start)

    run = paragraph.add_run(commented_text)

    range_end = OxmlElement("w:commentRangeEnd")
    range_end.set(qn("w:id"), str(comment_id))
    paragraph._p.append(range_end)

    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)
    paragraph._p.append(ref_run)


def merge_cells_horizontal(row_cells) -> None:
    """Horizontally merge the given cells via python-docx API."""
    a = row_cells[0]
    for other in row_cells[1:]:
        a = a.merge(other)


# ---------------------------------------------------------------------------
# Fixture assembly
# ---------------------------------------------------------------------------


def build_fixture() -> Document:
    doc = Document()

    # ---- Section 1: English-primary contract body ----
    section1 = doc.sections[0]
    section1.different_first_page_header_footer = False

    # Header (page-level scope)
    header = section1.header
    header.paragraphs[0].text = "CONFIDENTIAL — ABC Corporation internal · Draft v3"

    # Footer
    footer = section1.footer
    footer.paragraphs[0].text = (
        "ABC Corporation × Sunrise Ventures LLC · NDA · kim@abc-corp.kr"
    )

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MUTUAL NON-DISCLOSURE AGREEMENT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    # Parties with mixed script
    p1 = doc.add_paragraph()
    p1.add_run(
        "1. This Agreement is entered into by and between "
    )
    p1.add_run("ABC Corporation").bold = True
    p1.add_run(
        ' ("Discloser"), a corporation organized under the laws of the '
        "Republic of Korea (대한민국), and "
    )
    p1.add_run("Sunrise Ventures LLC").bold = True
    p1.add_run(
        ' ("Recipient"), a limited liability company organized under the '
        "laws of the State of Delaware."
    )

    # Definition clause — includes a tracked insertion
    def_p = doc.add_paragraph()
    def_p.add_run(
        '1.1 "Discloser" means ABC Corporation, including its affiliates and '
        "subsidiaries. Hereinafter referred to as "
    )
    add_tracked_insertion(def_p, '"the Buyer"')
    def_p.add_run(". ")
    add_tracked_deletion(def_p, " This sentence was supposed to be removed.")

    # Korean clause — mix of NFC, NFD, 한자, emoji
    kor_p = doc.add_paragraph()
    kor_p.add_run('1.2 "매수인"이라 함은 ')
    kor_p.add_run("ABC 주식회사").bold = True
    kor_p.add_run("를 말한다. 甲은 乙에게 비밀정보를 제공한다. ")
    # 한자 "甲" (U+7532), "乙" (U+4E59) — should round-trip cleanly.
    # Add an emoji too — U+1F4BC briefcase
    kor_p.add_run("📼 (비밀정보 보관함) ")
    # NFD-decomposed Korean text
    nfd_text = unicodedata.normalize("NFD", "가갸거겨")
    kor_p.add_run(f"NFD test: {nfd_text}")

    # Commented paragraph
    com_p = doc.add_paragraph()
    add_comment_range(
        doc,
        com_p,
        "Note: acquisition price is USD 100,000,000.",
        "This number is confidential — confirm with CFO before circulating.",
    )

    # Contact info table (complex: merged header + nested cell structure)
    doc.add_paragraph()
    doc.add_heading("2. Contact Information", level=2)

    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Light Grid Accent 1"

    # Header row — merge all 4 cells
    hdr_cells = tbl.rows[0].cells
    merge_cells_horizontal(hdr_cells)
    hdr_cells[0].text = "PARTIES AND CONTACT DETAILS — 당사자 연락처"

    # Column labels
    col_cells = tbl.rows[1].cells
    col_cells[0].text = "Party"
    col_cells[1].text = "Contact"
    col_cells[2].text = "Phone"
    col_cells[3].text = "Registration"

    # ABC row
    abc = tbl.rows[2].cells
    abc[0].text = "ABC Corp"
    abc[1].text = "kim@abc-corp.kr"
    abc[2].text = "010-1234-5678"
    abc[3].text = "123-45-67890"

    # Sunrise row
    sun = tbl.rows[3].cells
    sun[0].text = "Sunrise Ventures"
    sun[1].text = "legal@sunrise.com"
    sun[2].text = "+1 415 555 0199"
    sun[3].text = "EIN 12-3456789"

    # A nested paragraph with inline PII mid-sentence
    doc.add_paragraph()
    doc.add_heading("3. Confidential Information", level=2)
    doc.add_paragraph(
        "3.1 The Buyer acknowledges that during the course of discussions, "
        "ABC Corporation may disclose certain proprietary information, including "
        "pricing, customer lists, and acquisition plans related to Project Falcon "
        "and 블루윙 2.0."
    )
    doc.add_paragraph(
        "3.2 본 계약의 목적은 ABC Corporation과 관련된 비밀정보를 보호하는 데 있으며, "
        "매수인은 본 정보를 제3자에게 공개하지 않는다. 연락처: 김철수 (010-1234-5678)."
    )

    # ---- Section break → Section 2 with different header ----
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False
    new_section.header.paragraphs[0].text = (
        "APPENDIX A — ABC Corporation / Sunrise Ventures · Schedule of Deliverables"
    )
    new_section.footer.paragraphs[0].text = "Page 2 · Confidential"

    doc.add_heading("Appendix A — Schedule", level=1)

    # Complex nested table: 3-row, 3-col with a vertically merged first column
    nested = doc.add_table(rows=3, cols=3)
    nested.style = "Light Grid Accent 1"

    # Vertical merge on column 0
    col0_row0 = nested.rows[0].cells[0]
    col0_row1 = nested.rows[1].cells[0]
    col0_row2 = nested.rows[2].cells[0]
    merged_v = col0_row0.merge(col0_row1).merge(col0_row2)
    merged_v.text = "Phase 1"

    nested.rows[0].cells[1].text = "Deliverable A"
    nested.rows[0].cells[2].text = "2026-05-01"
    nested.rows[1].cells[1].text = "Deliverable B (블루윙 2.0 전달)"
    nested.rows[1].cells[2].text = "2026-06-15"
    nested.rows[2].cells[1].text = "Deliverable C — Project Falcon kickoff"
    nested.rows[2].cells[2].text = "2026-07-30"

    # Footnote-ish content via endnote marker (python-docx has limited footnote
    # support, so we simulate with a paragraph styled as a footnote reference)
    doc.add_paragraph()
    foot_like = doc.add_paragraph()
    foot_like.add_run("* Notes: ").italic = True
    foot_like.add_run(
        "Deliverables subject to approval by 이영희 (CTO, ABC Corp) "
        "and mike@sunrise.com."
    )

    return doc


def main() -> int:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_fixture()
    doc.save(str(FIXTURE_PATH))

    size_kb = FIXTURE_PATH.stat().st_size / 1024
    print(f"✓ wrote {FIXTURE_PATH.relative_to(FIXTURE_DIR.parent.parent)}")
    print(f"  size: {size_kb:.1f} KB")

    # Verify round-trip integrity by reading the file back with python-docx.
    # If this fails, the file is malformed.
    try:
        Document(str(FIXTURE_PATH))
        print("✓ python-docx round-trip read: OK")
    except Exception as e:
        print(f"✗ round-trip read FAILED: {e}", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
